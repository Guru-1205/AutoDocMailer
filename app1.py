from flask import Flask, render_template, request, send_file
from flask_mail import Mail, Message
from io import BytesIO
from docx import Document
from docx.shared import Pt
from datetime import datetime
import os
from dotenv import load_dotenv

load_dotenv()

app = Flask(__name__)

# Flask-Mail configuration
app.config['MAIL_SERVER'] = 'smtp.gmail.com'
app.config['MAIL_PORT'] = 465
app.config['MAIL_USE_SSL'] = True
app.config['MAIL_USERNAME'] = os.getenv('SENDER_MAIL_USERNAME')
app.config['MAIL_PASSWORD'] = os.getenv('SENDER_MAIL_APP_PASSWORD')

mail = Mail(app)

template_path = 'static/template.docx'
custom_font_path = 'static/fonts/times new roman bold.ttf'
custom_font_size = Pt(14)  # Set the desired font size


@app.route('/')
def index():
    return render_template('index.html')


@app.route('/process', methods=['POST'])
def process():
    # Get user input from the form
    replace_dict = {
        '{{activityName}}': request.form['name'],
        '{{activityFromDate}}': format_date(request.form['fromDate']),
        '{{activityYear}}': request.form['year'],
        '{{activityDate}}': format_date(request.form['date']),
        '{{activityPlace}}': request.form['place'],
        '{{activityDuration}}': request.form['timing'],
        '{{activityDay}}': request.form['day']
    }

    # Load the Word document template
    with open(template_path, 'rb') as doc_stream:
        doc = Document(BytesIO(doc_stream.read()))

    # Replace placeholders with user input
    for paragraph in doc.paragraphs:
        for key, value in replace_dict.items():
            if key in paragraph.text:
                paragraph.text = paragraph.text.replace(key, str(value))

    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            run.font.size = custom_font_size
    # Create an in-memory stream for the modified document
    output_stream = BytesIO()
    doc.save(output_stream)

    # Move the stream cursor to the beginning
    output_stream.seek(0)

    # Send the modified document via email
    send_email(request.form['receiver_email'], output_stream)

    return "Email sent successfully!"


def format_date(raw_date):
    # Convert the raw date string to a datetime object
    date_object = datetime.strptime(raw_date, '%Y-%m-%d')
    # Format the date in day/month/year format
    formatted_date = date_object.strftime('%d/%m/%Y')
    return formatted_date


def send_email(receiver_email, attachment_stream):
    em = Message(
        'XYZ ORGANISATION LETTER',
        sender=os.getenv('SENDER_MAIL_USERNAME'),
        recipients=[receiver_email]
    )
    em.body = 'Activity letter'
    em.attach('Activity_letter.docx', 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
              attachment_stream.read())
    mail.send(em)


if __name__ == '__main__':
    app.run(debug=True)
